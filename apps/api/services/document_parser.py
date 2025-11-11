"""
Document parser service for parsing PDF and DOCX files
"""

import io
import logging
from typing import Dict

try:
    from PyPDF2 import PdfReader  # type: ignore
except ImportError:
    PdfReader = None  # type: ignore

try:
    import pdfplumber  # type: ignore
except ImportError:
    pdfplumber = None  # type: ignore

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None  # type: ignore

logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """Parser for PDF and DOCX documents"""

    def parse_pdf(self, file_content: bytes) -> str:
        """
        Parse PDF file and extract text with better structure preservation

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text from the PDF

        Raises:
            DocumentParserError: If parsing fails
        """
        # Try pdfplumber first for better layout preservation
        if pdfplumber:
            try:
                pdf_file = io.BytesIO(file_content)
                text_pages = []

                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        # Extract text with layout preserved
                        text = page.extract_text(layout=True)
                        if text:
                            text_pages.append(text)

                full_text = '\n\n=== PAGE BREAK ===\n\n'.join(text_pages)

                if full_text.strip():
                    return full_text

            except Exception as e:
                logger.warning(f"pdfplumber failed, falling back to PyPDF2: {str(e)}")

        # Fallback to PyPDF2
        if not PdfReader:  # type: ignore
            raise DocumentParserError("PDF parsing libraries not installed")

        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_pages = []
            for page in reader.pages:  # type: ignore
                # Extract text with layout mode for better structure
                text = page.extract_text(extraction_mode="layout")
                if text:
                    # Clean up excessive whitespace while preserving structure
                    lines = text.split('\n')
                    cleaned_lines = []
                    for line in lines:
                        # Keep lines with content, preserve some spacing
                        cleaned_line = line.rstrip()
                        if cleaned_line.strip():
                            cleaned_lines.append(cleaned_line)

                    text_pages.append('\n'.join(cleaned_lines))

            full_text = '\n\n=== PAGE BREAK ===\n\n'.join(text_pages)

            if not full_text.strip():
                raise DocumentParserError("No text content found in PDF")

            return full_text

        except Exception as e:
            logger.error(f"Failed to parse PDF: {str(e)}")
            raise DocumentParserError(f"Failed to parse PDF: {str(e)}")

    def parse_docx(self, file_content: bytes) -> str:
        """
        Parse DOCX file and extract text

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text from the DOCX

        Raises:
            DocumentParserError: If parsing fails
        """
        if not Document:  # type: ignore
            raise DocumentParserError("python-docx is not installed")

        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)

            full_text = '\n'.join(paragraphs)
            if table_text:
                full_text += '\n\n' + '\n'.join(table_text)

            if not full_text.strip():
                raise DocumentParserError("No text content found in DOCX")

            return full_text

        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}")
            raise DocumentParserError(f"Failed to parse DOCX: {str(e)}")

    def parse_document(self, file_content: bytes, filename: str) -> Dict[str, str]:
        """
        Parse document based on file extension

        Args:
            file_content: Document file content as bytes
            filename: Name of the file to determine type

        Returns:
            Dictionary with 'text' key containing extracted text

        Raises:
            DocumentParserError: If parsing fails or unsupported format
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            text = self.parse_pdf(file_content)
            return {'text': text, 'type': 'pdf'}
        elif filename_lower.endswith('.docx'):
            text = self.parse_docx(file_content)
            return {'text': text, 'type': 'docx'}
        else:
            raise DocumentParserError(
                f"Unsupported file format. Only PDF and DOCX are supported. Got: {filename}"
            )

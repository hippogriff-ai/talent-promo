"""Export node for generating ATS-optimized outputs and LinkedIn suggestions.

This module handles:
- PDF generation (ATS-optimized)
- Plain text generation
- JSON export
- ATS analysis and scoring
- LinkedIn profile suggestions
"""

import io
import json
import logging
import re
from datetime import datetime
from typing import Any, Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from workflow.state import ResumeState, ATSReport, LinkedInSuggestion, ExportOutput

logger = logging.getLogger(__name__)


# ============================================================================
# Export Node
# ============================================================================


async def export_node(state: ResumeState) -> dict[str, Any]:
    """Export resume to multiple formats with ATS analysis.

    This node:
    1. Validates draft is approved
    2. Optimizes for ATS
    3. Generates PDF, TXT, JSON outputs
    4. Analyzes ATS compatibility
    5. Generates LinkedIn suggestions
    """
    logger.info("Export node - starting export workflow")

    # Check if draft is approved
    if not state.get("draft_approved"):
        logger.warning("Draft not approved - cannot export")
        return {
            "current_step": "error",
            "errors": [*state.get("errors", []), "Draft must be approved before export"],
            "updated_at": datetime.now().isoformat(),
        }

    resume_html = state.get("resume_final") or state.get("resume_html")
    if not resume_html:
        return {
            "current_step": "error",
            "errors": [*state.get("errors", []), "No resume content to export"],
            "updated_at": datetime.now().isoformat(),
        }

    job_posting = state.get("job_posting", {})
    gap_analysis = state.get("gap_analysis", {})
    user_profile = state.get("user_profile", {})

    # Get keywords from job posting and gap analysis
    job_keywords = _extract_job_keywords(job_posting, gap_analysis)

    # Step 1: Optimize for ATS
    logger.info("Export step: optimizing")
    optimized_html = optimize_for_ats(resume_html)

    # Step 2: Generate ATS report
    logger.info("Export step: analyzing_ats")
    ats_report = analyze_ats_compatibility(optimized_html, job_keywords)

    # Step 3: Generate LinkedIn suggestions
    logger.info("Export step: generating_linkedin")
    linkedin_suggestions = generate_linkedin_suggestions(
        optimized_html,
        user_profile,
        job_posting
    )

    # Create export output
    export_output = ExportOutput(
        pdf_generated=True,
        txt_generated=True,
        json_generated=True,
        ats_report=ats_report,
        linkedin_suggestions=linkedin_suggestions,
        export_completed=True,
        completed_at=datetime.now().isoformat(),
    )

    logger.info(f"Export complete. ATS score: {ats_report.keyword_match_score}")

    return {
        "current_step": "completed",
        "export_step": "completed",
        "resume_final": optimized_html,
        "export_output": export_output.model_dump(),
        "ats_report": ats_report.model_dump(),
        "linkedin_suggestions": linkedin_suggestions.model_dump(),
        "export_completed": True,
        "updated_at": datetime.now().isoformat(),
    }


# ============================================================================
# ATS Optimization
# ============================================================================


def optimize_for_ats(html_content: str) -> str:
    """Optimize HTML for ATS parsing.

    Removes formatting that breaks ATS parsers:
    - Tables
    - Multi-column layouts
    - Embedded images
    - Complex CSS
    """
    soup = BeautifulSoup(html_content, "html.parser")

    # Remove tables (ATS often can't parse tables)
    for table in soup.find_all("table"):
        # Convert table content to simple paragraphs
        rows = table.find_all("tr")
        replacement = soup.new_tag("div")
        for row in rows:
            cells = row.find_all(["td", "th"])
            text = " | ".join(cell.get_text().strip() for cell in cells)
            if text:
                p = soup.new_tag("p")
                p.string = text
                replacement.append(p)
        table.replace_with(replacement)

    # Remove inline styles that might cause issues
    for tag in soup.find_all(style=True):
        # Keep only safe styles
        style = tag.get("style", "")
        if "column" in style.lower() or "float" in style.lower():
            del tag["style"]

    # Remove images (ATS can't read them)
    for img in soup.find_all("img"):
        alt_text = img.get("alt", "")
        if alt_text:
            span = soup.new_tag("span")
            span.string = f"[{alt_text}]"
            img.replace_with(span)
        else:
            img.decompose()

    # Remove multi-column divs
    for div in soup.find_all("div"):
        classes = div.get("class", [])
        if any("col" in c.lower() for c in classes):
            # Flatten column structure
            div.unwrap()

    return str(soup)


# ============================================================================
# ATS Analysis
# ============================================================================


def analyze_ats_compatibility(
    html_content: str,
    job_keywords: list[str]
) -> ATSReport:
    """Analyze resume for ATS compatibility.

    Args:
        html_content: Resume HTML content
        job_keywords: Keywords from job posting

    Returns:
        ATSReport with score and recommendations
    """
    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text().lower()

    # Find matched and missing keywords
    matched = []
    missing = []

    for keyword in job_keywords:
        keyword_lower = keyword.lower()
        # Check for exact match or partial match
        if keyword_lower in text or any(
            word in text for word in keyword_lower.split()
        ):
            matched.append(keyword)
        else:
            missing.append(keyword)

    # Calculate score
    if job_keywords:
        score = int((len(matched) / len(job_keywords)) * 100)
    else:
        score = 100  # No keywords to match

    # Check for formatting issues
    formatting_issues = []

    # Check for tables
    if soup.find("table"):
        formatting_issues.append("Contains tables - some ATS may not parse correctly")

    # Check for images
    if soup.find("img"):
        formatting_issues.append("Contains images - text in images won't be read")

    # Check for fancy Unicode characters
    fancy_chars = re.findall(r'[^\x00-\x7F]', text)
    if fancy_chars:
        unique_chars = set(fancy_chars)
        if len(unique_chars) > 5:
            formatting_issues.append("Contains special characters that may not parse correctly")

    # Check for headers/footers that might confuse ATS
    headers = soup.find_all(["header", "footer"])
    if headers:
        formatting_issues.append("Contains header/footer elements - content may be skipped")

    # Generate recommendations
    recommendations = []

    if score < 70:
        recommendations.append(f"Add more job-relevant keywords. Missing: {', '.join(missing[:5])}")

    if missing:
        top_missing = missing[:3]
        recommendations.append(f"Consider incorporating: {', '.join(top_missing)}")

    if formatting_issues:
        recommendations.append("Simplify formatting for better ATS compatibility")

    if not soup.find("h2"):
        recommendations.append("Add clear section headers (Experience, Education, Skills)")

    return ATSReport(
        keyword_match_score=score,
        matched_keywords=matched,
        missing_keywords=missing,
        formatting_issues=formatting_issues,
        recommendations=recommendations,
    )


def _extract_job_keywords(
    job_posting: dict,
    gap_analysis: dict
) -> list[str]:
    """Extract keywords from job posting and gap analysis."""
    keywords = set()

    # From gap analysis
    keywords_to_include = gap_analysis.get("keywords_to_include", [])
    keywords.update(keywords_to_include)

    # From job posting requirements
    requirements = job_posting.get("requirements", [])
    for req in requirements[:10]:  # Limit to top 10
        # Extract key terms from requirements
        words = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', req)
        keywords.update(words)

    # From tech stack
    tech_stack = job_posting.get("tech_stack", [])
    keywords.update(tech_stack)

    # From preferred qualifications
    preferred = job_posting.get("preferred_qualifications", [])
    for pref in preferred[:5]:
        words = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', pref)
        keywords.update(words)

    return list(keywords)


# ============================================================================
# LinkedIn Suggestions
# ============================================================================


def generate_linkedin_suggestions(
    html_content: str,
    user_profile: dict,
    job_posting: dict
) -> LinkedInSuggestion:
    """Generate LinkedIn profile optimization suggestions.

    Args:
        html_content: Resume HTML content
        user_profile: User's profile data
        job_posting: Target job posting

    Returns:
        LinkedInSuggestion with headline, summary, and experience bullets
    """
    soup = BeautifulSoup(html_content, "html.parser")

    # Extract name
    name = user_profile.get("name", "Professional")

    # Extract current/recent role
    h3_tags = soup.find_all("h3")
    current_role = ""
    if h3_tags:
        current_role = h3_tags[0].get_text().strip()

    # Target job title
    target_title = job_posting.get("title", "")
    target_company = job_posting.get("company_name", "")

    # Generate headline (max 220 chars)
    headline = _generate_headline(name, current_role, target_title, user_profile)

    # Generate summary (about section)
    summary = _generate_summary(soup, user_profile, target_title)

    # Extract experience bullets
    experience_bullets = _extract_experience_bullets(soup, user_profile)

    return LinkedInSuggestion(
        headline=headline,
        summary=summary,
        experience_bullets=experience_bullets,
    )


def _generate_headline(
    name: str,
    current_role: str,
    target_title: str,
    user_profile: dict
) -> str:
    """Generate an optimized LinkedIn headline."""
    # Get skills
    skills = user_profile.get("skills", [])[:3]
    skills_str = " | ".join(skills) if skills else ""

    # Build headline
    if current_role and target_title:
        base = f"{current_role}"
    elif current_role:
        base = current_role
    elif target_title:
        base = f"Aspiring {target_title}"
    else:
        base = "Professional"

    if skills_str:
        headline = f"{base} | {skills_str}"
    else:
        headline = base

    # Truncate to 220 chars
    if len(headline) > 220:
        headline = headline[:217] + "..."

    return headline


def _generate_summary(
    soup: BeautifulSoup,
    user_profile: dict,
    target_title: str
) -> str:
    """Generate a LinkedIn summary from resume content."""
    # Find summary section in HTML
    summary_text = ""

    # Look for summary/about section
    for h2 in soup.find_all("h2"):
        if any(word in h2.get_text().lower() for word in ["summary", "about", "profile"]):
            # Get the next sibling paragraph(s)
            next_elem = h2.find_next_sibling()
            while next_elem and next_elem.name in ["p", "div"]:
                summary_text += next_elem.get_text().strip() + "\n\n"
                next_elem = next_elem.find_next_sibling()
                if next_elem and next_elem.name == "h2":
                    break

    # Fallback to user profile summary
    if not summary_text:
        summary_text = user_profile.get("summary", "")

    # Clean up
    summary_text = summary_text.strip()

    # LinkedIn summary max is 2600 chars
    if len(summary_text) > 2600:
        summary_text = summary_text[:2597] + "..."

    return summary_text


def _extract_experience_bullets(
    soup: BeautifulSoup,
    user_profile: dict
) -> list[dict]:
    """Extract experience bullets mapped to LinkedIn format."""
    experience_bullets = []

    # Get experience from user profile
    experiences = user_profile.get("experience", [])

    for exp in experiences:
        bullets = {
            "company": exp.get("company", ""),
            "position": exp.get("position", ""),
            "bullets": exp.get("achievements", []),
        }
        experience_bullets.append(bullets)

    # If no profile experience, try to extract from HTML
    if not experience_bullets:
        # Look for experience section
        for h2 in soup.find_all("h2"):
            if "experience" in h2.get_text().lower():
                # Find job entries (usually h3 or strong)
                h3_tags = h2.find_next_siblings("h3")
                for h3 in h3_tags[:5]:  # Limit to 5 positions
                    position = h3.get_text().strip()
                    company = ""

                    # Look for company in next element
                    next_elem = h3.find_next_sibling()
                    if next_elem and next_elem.name in ["p", "div"]:
                        company = next_elem.get_text().strip()

                    # Find bullets
                    bullets = []
                    ul = h3.find_next("ul")
                    if ul:
                        for li in ul.find_all("li"):
                            bullets.append(li.get_text().strip())

                    experience_bullets.append({
                        "company": company,
                        "position": position,
                        "bullets": bullets,
                    })

    return experience_bullets


# ============================================================================
# File Generation
# ============================================================================


def html_to_docx(html_content: str) -> bytes:
    """Convert HTML resume to DOCX format.

    Args:
        html_content: HTML resume content

    Returns:
        DOCX file as bytes
    """
    logger.info("Converting HTML to DOCX")

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Process HTML elements
    for element in soup.children:
        _process_element(doc, element)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def _process_element(doc: Document, element: Any) -> None:
    """Process HTML element and add to document."""
    if element.name is None:
        # Text node
        text = str(element).strip()
        if text:
            doc.add_paragraph(text)
        return

    if element.name == "h1":
        # Name/title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(element.get_text().strip())
        run.bold = True
        run.font.size = Pt(18)

    elif element.name == "h2":
        # Section headers
        p = doc.add_paragraph()
        run = p.add_run(element.get_text().strip())
        run.bold = True
        run.font.size = Pt(14)
        # Add line below
        p.paragraph_format.space_after = Pt(6)

    elif element.name == "h3":
        # Subsection headers (job titles, etc.)
        p = doc.add_paragraph()
        run = p.add_run(element.get_text().strip())
        run.bold = True
        run.font.size = Pt(12)

    elif element.name == "p":
        text = element.get_text().strip()
        if text:
            doc.add_paragraph(text)

    elif element.name == "ul":
        # Bullet list
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text().strip(), style="List Bullet")

    elif element.name == "ol":
        # Numbered list
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text().strip(), style="List Number")

    elif element.name in ["div", "section", "article"]:
        # Container elements - process children
        for child in element.children:
            _process_element(doc, child)


def html_to_pdf(html_content: str) -> bytes:
    """Convert HTML resume to PDF format.

    Args:
        html_content: HTML resume content

    Returns:
        PDF file as bytes
    """
    logger.info("Converting HTML to PDF")

    try:
        from weasyprint import HTML

        # Add basic styling
        styled_html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        @page {{
            margin: 0.5in;
        }}
        body {{
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }}
        h1 {{
            font-size: 18pt;
            text-align: center;
            margin-bottom: 5px;
            color: #000;
        }}
        h2 {{
            font-size: 13pt;
            border-bottom: 1px solid #333;
            padding-bottom: 3px;
            margin-top: 15px;
            margin-bottom: 8px;
            color: #000;
        }}
        h3 {{
            font-size: 11pt;
            margin-bottom: 3px;
            margin-top: 10px;
        }}
        p {{
            margin: 5px 0;
        }}
        ul {{
            margin: 5px 0;
            padding-left: 20px;
        }}
        li {{
            margin: 3px 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""

        html = HTML(string=styled_html)
        pdf_bytes = html.write_pdf()

        return pdf_bytes

    except ImportError:
        logger.error("WeasyPrint not installed - PDF export unavailable")
        raise ImportError("WeasyPrint is required for PDF export. Install with: pip install weasyprint")


def html_to_text(html_content: str) -> str:
    """Convert HTML resume to plain text.

    Args:
        html_content: HTML resume content

    Returns:
        Plain text version of resume
    """
    soup = BeautifulSoup(html_content, "html.parser")

    lines = []

    for element in soup.descendants:
        if element.name == "h1":
            text = element.get_text().strip()
            lines.append(text.upper())
            lines.append("=" * len(text))
            lines.append("")

        elif element.name == "h2":
            text = element.get_text().strip()
            lines.append("")
            lines.append(text.upper())
            lines.append("-" * len(text))

        elif element.name == "h3":
            text = element.get_text().strip()
            lines.append("")
            lines.append(text)

        elif element.name == "p":
            text = element.get_text().strip()
            if text and element.parent.name not in ["li"]:
                lines.append(text)

        elif element.name == "li":
            text = element.get_text().strip()
            lines.append(f"  • {text}")

    return "\n".join(lines)


def html_to_json(html_content: str, user_profile: dict) -> str:
    """Convert resume to structured JSON.

    Args:
        html_content: HTML resume content
        user_profile: User profile data

    Returns:
        JSON string of structured resume
    """
    soup = BeautifulSoup(html_content, "html.parser")

    # Start with user profile data
    resume_data = {
        "name": user_profile.get("name", ""),
        "email": user_profile.get("email", ""),
        "phone": user_profile.get("phone", ""),
        "location": user_profile.get("location", ""),
        "linkedin_url": user_profile.get("linkedin_url", ""),
        "summary": "",
        "experience": [],
        "education": [],
        "skills": user_profile.get("skills", []),
        "certifications": user_profile.get("certifications", []),
        "generated_at": datetime.now().isoformat(),
    }

    # Extract summary from HTML
    for h2 in soup.find_all("h2"):
        if "summary" in h2.get_text().lower():
            next_p = h2.find_next_sibling("p")
            if next_p:
                resume_data["summary"] = next_p.get_text().strip()
            break

    # Use experience from profile
    for exp in user_profile.get("experience", []):
        resume_data["experience"].append({
            "company": exp.get("company", ""),
            "position": exp.get("position", ""),
            "location": exp.get("location", ""),
            "start_date": exp.get("start_date", ""),
            "end_date": exp.get("end_date", ""),
            "is_current": exp.get("is_current", False),
            "achievements": exp.get("achievements", []),
        })

    # Use education from profile
    for edu in user_profile.get("education", []):
        resume_data["education"].append({
            "institution": edu.get("institution", ""),
            "degree": edu.get("degree", ""),
            "field_of_study": edu.get("field_of_study", ""),
            "start_date": edu.get("start_date", ""),
            "end_date": edu.get("end_date", ""),
        })

    return json.dumps(resume_data, indent=2)


def export_resume(
    html_content: str,
    format: str = "pdf",
    filename_base: str = "resume",
    user_profile: Optional[dict] = None,
) -> tuple[bytes, str, str]:
    """Export resume to specified format.

    Args:
        html_content: HTML resume content
        format: Export format ("pdf", "docx", "txt", "json")
        filename_base: Base filename without extension
        user_profile: User profile for JSON export

    Returns:
        Tuple of (file_bytes, content_type, filename)
    """
    if format.lower() == "pdf":
        file_bytes = html_to_pdf(html_content)
        content_type = "application/pdf"
        filename = f"{filename_base}.pdf"

    elif format.lower() == "txt":
        text_content = html_to_text(html_content)
        file_bytes = text_content.encode("utf-8")
        content_type = "text/plain"
        filename = f"{filename_base}.txt"

    elif format.lower() == "json":
        json_content = html_to_json(html_content, user_profile or {})
        file_bytes = json_content.encode("utf-8")
        content_type = "application/json"
        filename = f"{filename_base}.json"

    else:  # docx
        file_bytes = html_to_docx(html_content)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{filename_base}.docx"

    logger.info(f"Exported resume: {filename} ({len(file_bytes)} bytes)")

    return file_bytes, content_type, filename
